from docx import Document
from datetime import date


class MemoBuilder:

    def __init__(self):
        self.document = Document()

    def build_memo(self, rebuttals: list, output_path: str):

        # Header Section
        self.document.add_heading("SAMPLE LEGAL REPLY TO ANONYMIZED TAX NOTICE", level=1)
        self.document.add_paragraph("Student Submission – AI-Assisted Legal Response")
        self.document.add_paragraph("Course: Legal Writing & AI Applications")
        self.document.add_paragraph("Submitted By: [Student Name]")
        self.document.add_paragraph(f"Date: {date.today()}")
        self.document.add_paragraph("")

        # Notice Reference
        self.document.add_heading("REPLY TO NOTICE REF: TAX/CASE/2023/789", level=2)
        self.document.add_paragraph("To:")
        self.document.add_paragraph("[Designation]")
        self.document.add_paragraph("[Tax Department, City]")
        self.document.add_paragraph("[State]")
        self.document.add_paragraph("")

        self.document.add_paragraph("From:")
        self.document.add_paragraph("M/s [Company ABC]")
        self.document.add_paragraph("[Address Line 1]")
        self.document.add_paragraph("[City, State, PIN]")
        self.document.add_paragraph("GSTIN: [XXAAAAA1234X1XZ]")
        self.document.add_paragraph("")

        self.document.add_paragraph(f"Date: {date.today()}")
        self.document.add_paragraph("Subject: Reply to Legal Intimation under Section 74(5)")
        self.document.add_paragraph("")

        self.document.add_heading("POINT-WISE LEGAL RESPONSE", level=2)
        self.document.add_paragraph("")

        # Charge-wise responses
        for rebuttal in rebuttals:
            self.document.add_paragraph(rebuttal["analysis"])
            self.document.add_paragraph("")

        # AI Tools Section
        self.document.add_heading("AI TOOLS & PROMPTS USED", level=2)
        self.document.add_paragraph("1. Google Gemini API")
        self.document.add_paragraph("2. Structured legal reasoning prompts")
        self.document.add_paragraph("")

        self.document.add_paragraph(
            "*Note: This is a sample educational submission. "
            "All legal arguments are for academic purposes only.*"
        )

        self.document.save(output_path)
